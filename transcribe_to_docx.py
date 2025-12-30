import argparse
import logging
import os
import sys
from typing import List, Optional

from dotenv import load_dotenv
from docx import Document
from openai import OpenAI, OpenAIError

from app.services import file_service
from app.services.api_clients.openai_gpt4o import OpenAIGPT4oTranscriptionAPI


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Transcribe audio with gpt-4o-transcribe and save corrected DOCX files. "
            "If no file is specified, all .mp3 files in the resources folder are processed."
        )
    )
    parser.add_argument("audio", nargs="?", help="Path to the audio file to transcribe.")
    parser.add_argument(
        "--language",
        default=os.environ.get("DEFAULT_LANGUAGE", "auto"),
        help="Language code hint (default: auto).",
    )
    parser.add_argument(
        "--context",
        default="",
        help="Optional context prompt to improve transcription accuracy.",
    )
    parser.add_argument(
        "--context-file",
        help="Path to a text file with additional context prompt.",
    )
    parser.add_argument(
        "--correction-model",
        default=os.environ.get("OPENAI_CORRECTION_MODEL", "gpt-5-nano"),
        help="OpenAI model for correction (default: gpt-5-nano).",
    )
    parser.add_argument(
        "--no-correction",
        action="store_true",
        help="Skip the correction step and save the raw transcript.",
    )
    parser.add_argument(
        "--correction-chunk-words",
        type=int,
        default=1200,
        help="Max words per correction chunk (default: 1200).",
    )
    return parser.parse_args()


def build_context_prompt(context: str, context_file: Optional[str]) -> str:
    parts: List[str] = []
    if context:
        parts.append(context.strip())
    if context_file:
        try:
            with open(context_file, "r", encoding="utf-8") as handle:
                file_text = handle.read().strip()
        except OSError as exc:
            raise ValueError(f"Failed to read context file: {exc}") from exc
        if file_text:
            parts.append(file_text)
    return "\n\n".join(parts)


def split_text_for_correction(text: str, max_words: int) -> List[str]:
    normalized = text.replace("\r\n", "\n").strip()
    if not normalized:
        return []

    paragraphs = [paragraph.strip() for paragraph in normalized.split("\n\n") if paragraph.strip()]
    chunks: List[str] = []
    current: List[str] = []
    current_words = 0

    for paragraph in paragraphs:
        word_count = len(paragraph.split())
        if word_count > max_words:
            if current:
                chunks.append("\n\n".join(current))
                current = []
                current_words = 0
            words = paragraph.split()
            for start in range(0, len(words), max_words):
                chunks.append(" ".join(words[start:start + max_words]))
            continue

        if current_words + word_count > max_words and current:
            chunks.append("\n\n".join(current))
            current = [paragraph]
            current_words = word_count
        else:
            current.append(paragraph)
            current_words += word_count

    if current:
        chunks.append("\n\n".join(current))
    return chunks


def correct_transcript(
    text: str,
    client: OpenAI,
    model: str,
    max_words: int,
) -> str:
    chunks = split_text_for_correction(text, max_words=max_words)
    if not chunks:
        return text

    system_prompt = (
        "You are a precise editor for Hungarian speech transcripts. You pay special attention to hungarian Seventh Day Adventist religious text. "
        "Fix punctuation, casing, and obvious grammar "
        "errors while preserving the original meaning. Do not add or remove content. Keep names, numbers, "
        "and any speaker labels exactly as they appear. Return only the corrected text."
    )

    corrected_chunks: List[str] = []
    total_chunks = len(chunks)
    for index, chunk in enumerate(chunks, start=1):
        print(f"Correcting chunk {index}/{total_chunks}...", flush=True)
        user_prompt = f"Transcript:\n{chunk}"
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ]
        )
        content = response.choices[0].message.content or ""
        corrected_chunks.append(content.strip())

    corrected_text = "\n\n".join(chunk for chunk in corrected_chunks if chunk)
    return corrected_text if corrected_text else text


def save_docx(text: str, output_path: str) -> None:
    document = Document()
    paragraphs = [paragraph.strip() for paragraph in text.replace("\r\n", "\n").split("\n\n") if paragraph.strip()]
    if not paragraphs:
        document.add_paragraph("")
    else:
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
    document.save(output_path)


def get_project_root() -> str:
    return os.path.dirname(os.path.abspath(__file__))


def collect_audio_files(audio_arg: Optional[str], project_root: str) -> List[str]:
    if audio_arg:
        return [os.path.abspath(audio_arg)]

    resources_dir = os.path.join(project_root, "resources")
    if not os.path.isdir(resources_dir):
        raise ValueError(f"No audio file specified and resources folder not found: {resources_dir}")

    mp3_files: List[str] = []
    for filename in sorted(os.listdir(resources_dir)):
        if not filename.lower().endswith(".mp3"):
            continue
        path = os.path.abspath(os.path.join(resources_dir, filename))
        if os.path.isfile(path):
            mp3_files.append(path)

    if not mp3_files:
        raise ValueError(f"No .mp3 files found in resources folder: {resources_dir}")

    return mp3_files


def build_output_path(audio_path: str, output_dir: str) -> str:
    base_name = os.path.splitext(os.path.basename(audio_path))[0]
    return os.path.join(output_dir, f"{base_name}.docx")


def transcribe_one(
    audio_path: str,
    output_path: str,
    transcriber: OpenAIGPT4oTranscriptionAPI,
    context_prompt: str,
    language: str,
    correction_model: str,
    correction_client: Optional[OpenAI],
    correction_chunk_words: int,
    no_correction: bool,
) -> bool:
    if not os.path.isfile(audio_path):
        print(f"ERROR: Audio file not found: {audio_path}", file=sys.stderr)
        return False

    if not file_service.allowed_file(os.path.basename(audio_path)):
        allowed = ", ".join(sorted(file_service.ALLOWED_EXTENSIONS))
        print(
            f"WARNING: File extension not in allowed list ({allowed}). Attempting anyway.",
            file=sys.stderr,
        )

    def progress_callback(message: str, is_error: bool = False) -> None:
        stream = sys.stderr if is_error else sys.stdout
        print(message, file=stream, flush=True)

    transcription_text, detected_language = transcriber.transcribe(
        audio_file_path=audio_path,
        language_code=language,
        progress_callback=progress_callback,
        context_prompt=context_prompt,
        original_filename=os.path.basename(audio_path),
    )

    if not transcription_text:
        print(f"ERROR: Transcription failed for {audio_path}", file=sys.stderr)
        return False

    print(f"Transcription complete. Detected language: {detected_language}", flush=True)

    final_text = transcription_text
    if not no_correction:
        print(f"Applying correction with {correction_model}...", flush=True)
        try:
            if correction_client is None:
                raise OpenAIError("Correction client not initialized.")
            final_text = correct_transcript(
                transcription_text,
                correction_client,
                correction_model,
                max_words=correction_chunk_words,
            )
        except OpenAIError as exc:
            print(f"ERROR: Correction failed: {exc}", file=sys.stderr)
            print("Falling back to raw transcription.", file=sys.stderr)
            final_text = transcription_text

    save_docx(final_text, output_path)
    print(f"Saved DOCX: {output_path}", flush=True)
    return True


def main() -> int:
    load_dotenv()
    logging.basicConfig(level=logging.INFO, format="%(message)s")

    args = parse_args()
    project_root = get_project_root()

    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        print("ERROR: OPENAI_API_KEY is not set. Add it to your .env file.", file=sys.stderr)
        return 1

    try:
        audio_files = collect_audio_files(args.audio, project_root)
    except ValueError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    output_dir = os.path.join(project_root, "output")
    os.makedirs(output_dir, exist_ok=True)

    try:
        context_prompt = build_context_prompt(args.context, args.context_file)
    except ValueError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    transcriber = OpenAIGPT4oTranscriptionAPI(api_key)
    correction_client = None
    if not args.no_correction:
        correction_client = OpenAI(api_key=api_key)

    failures = 0
    for audio_path in audio_files:
        output_path = build_output_path(audio_path, output_dir)
        print(f"Processing: {audio_path}", flush=True)
        success = transcribe_one(
            audio_path=audio_path,
            output_path=output_path,
            transcriber=transcriber,
            context_prompt=context_prompt,
            language=args.language,
            correction_model=args.correction_model,
            correction_client=correction_client,
            correction_chunk_words=args.correction_chunk_words,
            no_correction=args.no_correction,
        )
        if not success:
            failures += 1

    if failures:
        print(f"Completed with {failures} failure(s).", file=sys.stderr)
        return 1

    print("All files completed successfully.", flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
